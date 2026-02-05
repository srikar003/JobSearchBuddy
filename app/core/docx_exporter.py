from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


SECTION_HEADERS = {
    "PROFESSIONAL SUMMARY",
    "SKILLS",
    "WORK EXPERIENCE",
    "ACADEMIC EDUCATION",
    "ACADEMIC PROJECTS",
    "ACHIEVEMENTS",
    "PROJECTS",
    "CERTIFICATIONS",
}


def _set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


class DocxExporter:
    """
    ATS-safe but richer DOCX:
      - Times New Roman 12
      - Name centered, larger
      - Section headers bold
      - Company/title bold, dates italic
      - Bullets using Word list style
    """
    def export(self, resume_text: str, output_path: str | Path) -> None:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Set Normal style
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

        lines = [ln.rstrip() for ln in resume_text.splitlines()]
        line_index = 0
        current_section = None

        # Helper to add a normal paragraph
        def add_para(text: str, bold=False, italic=False, size=12, center=False, space_before=0, space_after=0):
            p = doc.add_paragraph()
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(space_before)
            p.space_after = Pt(space_after)
            r = p.add_run(text)
            _set_run_font(r, size=size, bold=bold, italic=italic)
            return p

        while line_index < len(lines):
            raw = lines[line_index]
            line = raw.strip()

            if not line:
                doc.add_paragraph("")
                line_index += 1
                continue

            # Name (first non-empty line)
            if line_index == 0:
                add_para(line, bold=True, size=16, center=True, space_after=2)
                line_index += 1
                continue

            # Contact (second line usually)
            if line_index == 1 and ("@" in line or "|" in line or "linkedin" in line.lower()):
                add_para(line, size=11, center=True, space_after=6)
                line_index += 1
                continue

            # Section headers
            if line.upper() in SECTION_HEADERS:
                current_section = line.upper()
                add_para(current_section, bold=True, size=13, space_before=10, space_after=4)
                line_index += 1
                continue

            # Work experience 3-line blocks
            if current_section == "WORK EXPERIENCE":
                next1 = lines[line_index + 1].strip() if line_index + 1 < len(lines) else ""
                next2 = lines[line_index + 2].strip() if line_index + 2 < len(lines) else ""

                looks_like_company = ("|" in line) or line.isupper()
                looks_like_dates = ("-" in next1) and any(
                    m in next1.upper()
                    for m in ["JAN", "FEB", "MAR", "APR", "MAY", "JUN", "JUL", "AUG", "SEP", "SEPT", "OCT", "NOV", "DEC", "PRESENT"]
                )
                looks_like_title = next2.isupper() and len(next2) <= 50

                if looks_like_company and looks_like_dates and looks_like_title:
                    add_para(line, bold=True, space_after=0)
                    add_para(next1, italic=True, size=11, space_after=0)
                    add_para(next2, bold=True, space_after=2)
                    line_index += 3
                    continue

            # Bullets
            if line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(line[2:].strip())
                _set_run_font(r)
                line_index += 1
                continue

            if line.startswith("•"):
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(line.lstrip("•").strip())
                _set_run_font(r)
                line_index += 1
                continue

            # Default
            add_para(line, space_after=0)
            line_index += 1

        doc.save(str(output_path))
